"""Build the CSR-format English working draft v3 docx from the markdown source.

Usage: python tools/build_english_manuscript_v3.py

Reads  manuscript/BabyRuleBlind_resubmission_draft.md
Writes manuscript/BabyRuleBlind_resubmission_working_draft_v3.docx

Formatting: US Letter, 1-inch margins, Times New Roman, footer page numbers,
keep-with-next headings/captions (no orphaned section titles), numbered table
captions with notes, anonymized core properties.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "manuscript" / "BabyRuleBlind_resubmission_draft.md"
OUT = ROOT / "manuscript" / "BabyRuleBlind_resubmission_working_draft_v3.docx"

BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(12)
TABLE_SIZE = Pt(9)
NOTE_SIZE = Pt(10)

INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")


def set_font(run, size=BODY_SIZE, bold=False, italic=False, name=BODY_FONT):
    run.font.name = name
    run.font.size = size
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        fonts.set(qn(attr), name)


def add_runs(par, text, size=BODY_SIZE, base_bold=False, base_italic=False):
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = par.add_run(part[2:-2])
            set_font(run, size, bold=True, italic=base_italic)
        elif part.startswith("`") and part.endswith("`"):
            run = par.add_run(part[1:-1])
            set_font(run, size, bold=base_bold, italic=base_italic, name="Consolas")
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = par.add_run(part[1:-1])
            set_font(run, size, bold=base_bold, italic=True)
        else:
            run = par.add_run(part)
            set_font(run, size, bold=base_bold, italic=base_italic)


def add_paragraph(doc, text, size=BODY_SIZE, space_after=6, keep_next=False,
                  align=None, base_bold=False, base_italic=False):
    par = doc.add_paragraph()
    fmt = par.paragraph_format
    fmt.space_after = Pt(space_after)
    fmt.space_before = Pt(0)
    fmt.keep_with_next = keep_next
    if align is not None:
        par.alignment = align
    add_runs(par, text, size=size, base_bold=base_bold, base_italic=base_italic)
    return par


def add_bullet(doc, text, size=BODY_SIZE):
    par = doc.add_paragraph(style="List Bullet")
    par.paragraph_format.space_after = Pt(3)
    add_runs(par, text, size=size)
    for run in par.runs:
        pass
    return par


def set_cell_borders(cell, top=None, bottom=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, spec in (("top", top), ("bottom", bottom)):
        if spec is None:
            continue
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), spec)
        el.set(qn("w:color"), "000000")


def add_table(doc, header, rows):
    ncols = len(header)
    table = doc.add_table(rows=len(rows) + 1, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "autofit")
    tbl_pr.append(layout)

    def fill(cell, text, bold=False, align_right=False):
        par = cell.paragraphs[0]
        par.paragraph_format.space_after = Pt(2)
        par.paragraph_format.space_before = Pt(2)
        if align_right:
            par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_runs(par, text, size=TABLE_SIZE, base_bold=bold)

    for j, text in enumerate(header):
        cell = table.rows[0].cells[j]
        fill(cell, text, bold=True)
        set_cell_borders(cell, top="12", bottom="6")
    # repeat the header row when a long table splits across pages
    header_tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    header_tr_pr.append(tbl_header)
    for i, row in enumerate(rows, start=1):
        last = i == len(rows)
        for j, text in enumerate(row):
            cell = table.rows[i].cells[j]
            numeric = j > 0 and bool(re.match(r"^[+−\-.\d\[]", text.strip()))
            fill(cell, text, align_right=numeric)
            if last:
                set_cell_borders(cell, bottom="12")
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant = OxmlElement("w:cantSplit")
        tr_pr.append(cant)
    # keep the whole table on one page: every paragraph except those in the
    # final row keeps with next
    for row in table.rows[:-1]:
        for cell in row.cells:
            for par in cell.paragraphs:
                par.paragraph_format.keep_with_next = True
    return table


def add_page_number_footer(section):
    footer = section.footer
    par = footer.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run = par.add_run()
    set_font(run, Pt(10))
    run._element.append(fld_begin)
    run2 = par.add_run()
    set_font(run2, Pt(10))
    run2._element.append(instr)
    run3 = par.add_run()
    set_font(run3, Pt(10))
    run3._element.append(fld_end)


def new_document() -> Document:
    """Blank US-Letter TNR document with footer page numbers."""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for side in ("left", "right", "top", "bottom"):
        setattr(section, f"{side}_margin", Inches(1))
    add_page_number_footer(section)
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE
    return doc


def render_markdown(doc: Document, lines: list[str]) -> None:
    """Render manuscript-flavored markdown lines into doc."""
    i = 0
    n = len(lines)
    first_heading_seen = False
    while i < n:
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        if line.startswith("# ") and not first_heading_seen:
            add_paragraph(doc, line[2:], size=Pt(16), base_bold=True,
                          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10, keep_next=True)
            first_heading_seen = True
            i += 1
            # author line follows
            while i < n and not lines[i].strip():
                i += 1
            if i < n and not lines[i].startswith("#"):
                add_paragraph(doc, lines[i].strip(), size=Pt(12),
                              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
                i += 1
            continue
        if line.startswith("## "):
            add_paragraph(doc, line[3:], size=Pt(13), base_bold=True,
                          space_after=6, keep_next=True)
            i += 1
            continue
        if line.startswith("### "):
            add_paragraph(doc, line[4:], size=Pt(12), base_bold=True,
                          space_after=4, keep_next=True)
            i += 1
            continue
        if line.startswith("|"):
            block = []
            while i < n and lines[i].startswith("|"):
                block.append(lines[i])
                i += 1
            cells = [
                [c.strip() for c in row.strip().strip("|").split("|")]
                for row in block
            ]
            header = cells[0]
            body = [r for r in cells[1:] if not all(set(c) <= {"-", ":", " "} or c == "" for c in r)]
            add_table(doc, header, body)
            add_paragraph(doc, "", size=Pt(4), space_after=2)
            continue
        if line.startswith("- "):
            add_bullet(doc, line[2:])
            i += 1
            continue
        if re.match(r"^\d+\. ", line):
            par = doc.add_paragraph(style="List Number")
            par.paragraph_format.space_after = Pt(3)
            add_runs(par, re.sub(r"^\d+\. ", "", line))
            i += 1
            continue
        keep = line.startswith("**Table")
        note = line.startswith("*Note.*")
        add_paragraph(doc, line, size=NOTE_SIZE if note else BODY_SIZE,
                      keep_next=keep, space_after=4 if (keep or note) else 6)
        i += 1


def anonymize_core_properties(doc: Document, title: str) -> None:
    core = doc.core_properties
    core.author = ""
    core.last_modified_by = ""
    core.comments = ""
    core.title = title
    core.category = ""
    core.subject = ""


MANUSCRIPT_TITLE = ("Binary Outcome-Feedback Supervision of Memoryless LLM "
                    "Outputs under a Budgeted Intervention Policy")


def main() -> None:
    lines = SRC.read_text(encoding="utf-8").splitlines()
    doc = new_document()
    render_markdown(doc, lines)
    anonymize_core_properties(doc, MANUSCRIPT_TITLE)
    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
