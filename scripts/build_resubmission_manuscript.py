"""Build the resubmission working manuscript from its Markdown source.

The script starts from the prior anonymous Word document so its page setup,
theme, and heading styles remain the formatting authority. All prior body
content is removed before the current, evidence-checked manuscript is added.
It also consolidates the current verified bibliography with a small set of
still-used references from the withdrawn manuscript.
"""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.opc.constants import RELATIONSHIP_TYPE as RT


ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = ROOT / "manuscript" / "BabyRuleBlind_resubmission_draft.md"
OUTPUT_DOCX = ROOT / "manuscript" / "BabyRuleBlind_resubmission_working_draft.docx"
OUTPUT_BIB = ROOT / "manuscript" / "REFERENCES.bib"

PRIOR_PROJECT = ROOT.parent / "brain like"
TEMPLATE_DOCX = PRIOR_PROJECT / "TMLR" / "Baby40_anonymous.docx"
CURRENT_BIB = ROOT / "REFERENCES.bib"
WITHDRAWN_BIB = (
    PRIOR_PROJECT
    / "Cognitive Systems Research"
    / "최종구조"
    / "files"
    / "baby40_refs_updated.bib"
)

OLD_BIB_KEYS = [
    "Grant1948",
    "Kopp2021",
    "Miles2021",
    "Milner1963",
    "Binz2023",
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=80, bottom=70, end=80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_inches: list[float]) -> None:
    """Set one authoritative fixed-width geometry for table, grid, and cells."""
    widths_twips = [round(width * 1440) for width in widths_inches]
    total_twips = sum(widths_twips)
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    # Word draws the outer border at table indent minus the start cell margin.
    # Matching the 80-DXA start margin keeps that border aligned to body text.
    tbl_ind.set(qn("w:w"), "80")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid_cols = table._tbl.tblGrid.gridCol_lst
    for idx, width_twips in enumerate(widths_twips):
        grid_cols[idx].set(qn("w:w"), str(width_twips))

    for row in table.rows:
        for idx, width_twips in enumerate(widths_twips):
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width_twips))
            tc_w.set(qn("w:type"), "dxa")


def clear_document_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)
    # The copied template carried six old result figures. Removing the body
    # does not automatically drop their package relationships, so explicitly
    # remove them to prevent obsolete results from remaining hidden in the DOCX.
    for rel_id, rel in list(document.part.rels.items()):
        if rel.reltype == RT.IMAGE:
            document.part.drop_rel(rel_id)


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.space_after = Pt(6)

    for name, size, before, after in (
        ("Heading 1", 14, 12, 5),
        ("Heading 2", 12, 10, 4),
        ("Heading 3", 11, 8, 3),
    ):
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    if "Manuscript Title" not in styles:
        title_style = styles.add_style("Manuscript Title", WD_STYLE_TYPE.PARAGRAPH)
    else:
        title_style = styles["Manuscript Title"]
    title_style.font.name = "Times New Roman"
    title_style.font.size = Pt(16)
    title_style.font.bold = True
    title_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)
    title_style.paragraph_format.keep_with_next = True

    if "Abstract Text" not in styles:
        abstract_style = styles.add_style("Abstract Text", WD_STYLE_TYPE.PARAGRAPH)
    else:
        abstract_style = styles["Abstract Text"]
    abstract_style.font.name = "Times New Roman"
    abstract_style.font.size = Pt(10)
    abstract_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    abstract_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abstract_style.paragraph_format.left_indent = Inches(0.25)
    abstract_style.paragraph_format.right_indent = Inches(0.25)
    abstract_style.paragraph_format.space_after = Pt(6)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_end])


def set_page_geometry(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.4)
        section.footer_distance = Inches(0.4)
        if not section.footer.paragraphs:
            footer_p = section.footer.add_paragraph()
        else:
            footer_p = section.footer.paragraphs[0]
        footer_p.clear()
        add_page_number(footer_p)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows


def add_markdown_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    n_cols = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"

    available = 6.5
    header_lengths = [max(5, len(rows[0][i]) if i < len(rows[0]) else 5) for i in range(n_cols)]
    total = sum(header_lengths)
    widths = [max(0.55, min(1.55, available * x / total)) for x in header_lengths]
    scale = available / sum(widths)
    widths = [x * scale for x in widths]
    set_table_geometry(table, widths)

    for r_idx, row in enumerate(rows):
        for c_idx in range(n_cols):
            cell = table.cell(r_idx, c_idx)
            cell.width = Inches(widths[c_idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            text = row[c_idx] if c_idx < len(row) else ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(8 if n_cols >= 6 else 9)
            if r_idx == 0:
                run.bold = True
                set_cell_shading(cell, "D9E2F3")
        if r_idx == 0:
            set_repeat_table_header(table.rows[0])

    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_list_item(document: Document, text: str, ordered: bool = False) -> None:
    style = "List Number" if ordered else "List Bullet"
    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run(text)


def add_markdown(document: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    idx = 0
    first_title = True
    abstract_mode = False
    references_mode = False
    author_consumed = False

    while idx < len(lines):
        raw = lines[idx].rstrip()
        stripped = raw.strip()
        if not stripped:
            idx += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                table_lines.append(lines[idx].strip())
                idx += 1
            add_markdown_table(document, parse_table(table_lines))
            continue

        if stripped.startswith("# "):
            text = stripped[2:].strip()
            if first_title:
                p = document.add_paragraph(style="Manuscript Title")
                p.add_run(text)
                first_title = False
            else:
                document.add_paragraph(text, style="Heading 1")
            idx += 1
            continue

        if stripped.startswith("## "):
            text = stripped[3:].strip()
            abstract_mode = text == "Abstract"
            references_mode = text == "References"
            if references_mode:
                p = document.add_paragraph()
                p.add_run().add_break(WD_BREAK.PAGE)
            document.add_paragraph(text, style="Heading 1")
            idx += 1
            continue

        if stripped.startswith("### "):
            document.add_paragraph(stripped[4:].strip(), style="Heading 2")
            abstract_mode = False
            idx += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            add_list_item(document, re.sub(r"^\d+\.\s+", "", stripped), ordered=True)
            idx += 1
            continue

        if stripped.startswith("- "):
            add_list_item(document, stripped[2:])
            idx += 1
            continue

        if not author_consumed and not first_title:
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(stripped)
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            author_consumed = True
            idx += 1
            continue

        style = "Abstract Text" if abstract_mode else "Normal"
        paragraph = document.add_paragraph(style=style)
        if references_mode:
            paragraph.paragraph_format.left_indent = Inches(0.3)
            paragraph.paragraph_format.first_line_indent = Inches(-0.3)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.keep_together = True
        paragraph.add_run(stripped)
        idx += 1


def extract_bib_entries(text: str) -> dict[str, str]:
    entries: dict[str, str] = {}
    starts = list(re.finditer(r"(?m)^@([A-Za-z]+)\s*\{\s*([^,\s]+)\s*,", text))
    for i, match in enumerate(starts):
        start = match.start()
        end = starts[i + 1].start() if i + 1 < len(starts) else len(text)
        key = match.group(2)
        entries[key] = text[start:end].strip()
    return entries


def build_bibliography() -> None:
    current_text = CURRENT_BIB.read_text(encoding="utf-8-sig")
    old_text = WITHDRAWN_BIB.read_text(encoding="utf-8-sig")
    current_entries = extract_bib_entries(current_text)
    old_entries = extract_bib_entries(old_text)
    missing = [key for key in OLD_BIB_KEYS if key not in old_entries]
    if missing:
        raise KeyError(f"Missing selected withdrawn-manuscript BibTeX keys: {missing}")

    parts = [
        "% Resubmission working bibliography.",
        "% Current verified entries are followed by selected still-cited entries",
        "% recovered from the withdrawn manuscript bibliography.",
        "",
    ]
    parts.extend(current_entries[key] for key in current_entries)
    parts.append("")
    parts.append("% Selected references retained from the withdrawn manuscript.")
    parts.extend(old_entries[key] for key in OLD_BIB_KEYS)
    OUTPUT_BIB.write_text("\n\n".join(parts).rstrip() + "\n", encoding="utf-8")


def build_document() -> None:
    if not TEMPLATE_DOCX.exists():
        raise FileNotFoundError(TEMPLATE_DOCX)
    shutil.copy2(TEMPLATE_DOCX, OUTPUT_DOCX)
    document = Document(OUTPUT_DOCX)
    clear_document_body(document)
    configure_styles(document)
    set_page_geometry(document)
    add_markdown(document, SOURCE_MD.read_text(encoding="utf-8-sig"))

    props = document.core_properties
    props.title = "Sparse Outcome-Feedback Supervision of Memoryless LLM Outputs in a Rule-Shift Task"
    props.subject = "Anonymous resubmission working draft"
    props.author = ""
    props.last_modified_by = ""
    props.comments = ""
    props.keywords = "LLM; rule shifting; outcome feedback; external supervision"
    props.category = "Research manuscript"
    document.save(OUTPUT_DOCX)


if __name__ == "__main__":
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    build_bibliography()
    build_document()
    print(OUTPUT_DOCX)
    print(OUTPUT_BIB)
